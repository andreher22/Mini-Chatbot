"""
Script para generar la documentación formal del chatbot LóngBot.
Formato Word (.docx) siguiendo normas IEEE 830 y estilo APA.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ============================================================
# Configurar estilos APA
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 2.0  # Doble espacio (APA)

# Configurar márgenes (2.54 cm = 1 pulgada)
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_heading_apa(text, level=1):
    """Agregar encabezado con formato APA."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph_bold(text):
    """Agregar párrafo con texto en negrita."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_normal_paragraph(text):
    """Agregar párrafo normal."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.27)  # Sangría APA
    return p

# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LóngBot: Sistema de Chatbot Inteligente')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Técnica del Sistema')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Documento de Especificación de Requisitos de Software (SRS)\n').font.name = 'Times New Roman'
p.add_run('Basado en IEEE 830-1998 / ISO/IEC 25010\n').font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Proyecto: Mini-Chatbot con Django\n')
run.font.name = 'Times New Roman'
p.add_run('Versión: 2.0\n').font.name = 'Times New Roman'
p.add_run('Fecha: Marzo 2026\n').font.name = 'Times New Roman'
p.add_run('Materia: Sistemas Inteligentes\n').font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDOS
# ============================================================
add_heading_apa('Tabla de Contenidos', level=1)
contenidos = [
    '1. Introducción',
    '   1.1. Propósito',
    '   1.2. Alcance',
    '   1.3. Definiciones y Acrónimos',
    '   1.4. Referencias',
    '2. Descripción General',
    '   2.1. Perspectiva del Producto',
    '   2.2. Funcionalidades del Producto',
    '   2.3. Características de los Usuarios',
    '   2.4. Restricciones',
    '3. Requisitos Específicos',
    '   3.1. Requisitos Funcionales',
    '   3.2. Requisitos No Funcionales',
    '4. Arquitectura del Sistema',
    '   4.1. Arquitectura en Capas',
    '   4.2. Red Neuronal (MLPClassifier)',
    '   4.3. Módulo Calculadora',
    '5. Modelo de Datos',
    '6. Interfaz de Usuario',
    '7. Plan de Pruebas',
    '8. Ubicación de Funciones Principales',
    '9. Referencias Bibliográficas',
]
for item in contenidos:
    p = doc.add_paragraph(item)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
add_heading_apa('1. Introducción', level=1)

add_heading_apa('1.1. Propósito', level=2)
add_normal_paragraph(
    'El presente documento tiene como propósito describir de manera formal y detallada '
    'el sistema de chatbot inteligente denominado "LóngBot". Este documento sigue las '
    'directrices establecidas por el estándar IEEE 830-1998 para la Especificación de '
    'Requisitos de Software (SRS) y las normas de calidad ISO/IEC 25010 para evaluar '
    'las características de calidad del software.'
)
add_normal_paragraph(
    'La documentación está dirigida a desarrolladores, evaluadores, docentes y cualquier '
    'persona interesada en comprender la arquitectura, funcionalidades y modelo de '
    'inteligencia artificial implementado en el sistema.'
)

add_heading_apa('1.2. Alcance', level=2)
add_normal_paragraph(
    'LóngBot es un sistema de chatbot web desarrollado con Django 5.2.12 que utiliza '
    'una red neuronal artificial (MLPClassifier de scikit-learn) para detectar la intención '
    'del usuario a partir de mensajes de texto en español. El sistema es capaz de reconocer '
    '16 intenciones diferentes y además incorpora una calculadora integrada para resolver '
    'operaciones aritméticas básicas (suma, resta, multiplicación y división).'
)
add_normal_paragraph(
    'El sistema NO incluye: conexión a APIs externas, procesamiento de voz, soporte '
    'multiidioma, ni almacenamiento de sesiones de usuario individual.'
)

add_heading_apa('1.3. Definiciones y Acrónimos', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Término'
hdr[1].text = 'Definición'
definitions = [
    ('NLP', 'Procesamiento de Lenguaje Natural (Natural Language Processing)'),
    ('MLP', 'Perceptrón Multicapa (Multi-Layer Perceptron)'),
    ('TF-IDF', 'Frecuencia de Término - Frecuencia Inversa de Documento'),
    ('ReLU', 'Rectified Linear Unit, función de activación'),
    ('AJAX', 'Asynchronous JavaScript and XML'),
    ('SRS', 'Especificación de Requisitos de Software'),
    ('IEEE', 'Instituto de Ingenieros Eléctricos y Electrónicos'),
    ('ISO', 'Organización Internacional de Normalización'),
    ('API', 'Interfaz de Programación de Aplicaciones'),
    ('Singleton', 'Patrón de diseño que garantiza una sola instancia de una clase'),
    ('Data Augmentation', 'Técnica de multiplicación de datos de entrenamiento'),
]
for term, defn in definitions:
    row = table.add_row().cells
    row[0].text = term
    row[1].text = defn

add_heading_apa('1.4. Referencias', level=2)
refs = [
    'IEEE. (1998). IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications.',
    'ISO/IEC. (2011). ISO/IEC 25010:2011, Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE).',
    'Scikit-learn. (2024). MLPClassifier — Multi-layer Perceptron classifier. https://scikit-learn.org/stable/modules/generated/sklearn.neural_network.MLPClassifier.html',
    'Django Software Foundation. (2024). Django Documentation (v5.2). https://docs.djangoproject.com/en/5.2/',
    'Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.',
]
for ref in refs:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

# ============================================================
# 2. DESCRIPCIÓN GENERAL
# ============================================================
add_heading_apa('2. Descripción General', level=1)

add_heading_apa('2.1. Perspectiva del Producto', level=2)
add_normal_paragraph(
    'LóngBot es un producto de software independiente desarrollado como proyecto académico '
    'para la materia de Sistemas Inteligentes. El sistema combina técnicas de procesamiento '
    'de lenguaje natural con una interfaz web elegante inspirada en la filosofía china, '
    'creando una experiencia de usuario única y educativa.'
)
add_normal_paragraph(
    'El sistema opera como una aplicación web de página única (SPA) que se comunica '
    'con el backend de Django mediante peticiones AJAX. No requiere servidores externos '
    'ni APIs de terceros para su funcionamiento básico.'
)

add_heading_apa('2.2. Funcionalidades del Producto', level=2)
funcionalidades = [
    ('RF-01', 'Detección de Intenciones', 'El sistema detecta 16 intenciones diferentes del usuario mediante una red neuronal MLP.'),
    ('RF-02', 'Respuestas Inteligentes', 'Genera respuestas contextuales basadas en la intención detectada, con variedad aleatoria.'),
    ('RF-03', 'Calculadora Integrada', 'Resuelve operaciones de suma, resta, multiplicación y división con detección por regex.'),
    ('RF-04', 'Hora y Fecha Dinámica', 'Genera respuestas dinámicas con la hora y fecha actual del servidor.'),
    ('RF-05', 'Historial de Conversaciones', 'Almacena todas las conversaciones en base de datos SQLite para consulta posterior.'),
    ('RF-06', 'Interfaz Interactiva', 'Frontend con partículas, animaciones, barra de confianza y visualización de red neuronal.'),
    ('RF-07', 'Respuesta por Defecto', 'Manejo elegante de mensajes no reconocidos con citas filosóficas chinas.'),
    ('RF-08', 'Panel de Administración', 'Acceso al panel Django Admin para consultar el historial de conversaciones.'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = 'Funcionalidad'
hdr[2].text = 'Descripción'
for fid, fname, fdesc in funcionalidades:
    row = table.add_row().cells
    row[0].text = fid
    row[1].text = fname
    row[2].text = fdesc

add_heading_apa('2.3. Características de los Usuarios', level=2)
add_normal_paragraph(
    'El sistema está diseñado para dos tipos de usuarios: (1) Usuarios finales que interactúan '
    'con el chatbot a través de la interfaz web para obtener información, realizar cálculos '
    'y explorar filosofía china; y (2) Administradores que acceden al panel de Django para '
    'consultar el historial de conversaciones y monitorear el rendimiento del sistema.'
)

add_heading_apa('2.4. Restricciones', level=2)
restricciones = [
    'El sistema opera únicamente en español.',
    'Requiere Python 3.10 o superior.',
    'La red neuronal se entrena en memoria al iniciar el servidor (no persiste en disco).',
    'La calculadora solo soporta operaciones aritméticas básicas con dos operandos.',
    'No incluye autenticación de usuarios finales.',
]
for r in restricciones:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. REQUISITOS ESPECÍFICOS
# ============================================================
add_heading_apa('3. Requisitos Específicos', level=1)

add_heading_apa('3.1. Requisitos Funcionales', level=2)
req_func = [
    ('RF-01', 'Envío de Mensajes', 'El usuario debe poder enviar mensajes de texto a través de la interfaz web.', 'Alta'),
    ('RF-02', 'Detección de Intenciones', 'El sistema debe detectar la intención del usuario con una confianza mínima del 35%.', 'Alta'),
    ('RF-03', 'Calculadora', 'El sistema debe resolver operaciones aritméticas (+, -, *, /) escritas en formato natural o simbólico.', 'Alta'),
    ('RF-04', 'Respuesta Dinámica', 'El sistema debe generar respuestas dinámicas para hora y fecha actual.', 'Media'),
    ('RF-05', 'Persistencia', 'Todas las conversaciones deben almacenarse en la base de datos.', 'Media'),
    ('RF-06', 'Fallback Inteligente', 'Mensajes no reconocidos deben recibir una respuesta por defecto con citas filosóficas.', 'Media'),
    ('RF-07', 'Visualización NN', 'La interfaz debe mostrar la arquitectura de la red neuronal.', 'Baja'),
    ('RF-08', 'Panel Admin', 'El administrador debe poder consultar el historial desde Django Admin.', 'Baja'),
]
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = 'Requisito'
hdr[2].text = 'Descripción'
hdr[3].text = 'Prioridad'
for rid, rname, rdesc, rpri in req_func:
    row = table.add_row().cells
    row[0].text = rid
    row[1].text = rname
    row[2].text = rdesc
    row[3].text = rpri

add_heading_apa('3.2. Requisitos No Funcionales (ISO/IEC 25010)', level=2)
req_nofunc = [
    ('RNF-01', 'Rendimiento', 'El tiempo de respuesta del chatbot no debe exceder 2 segundos.'),
    ('RNF-02', 'Usabilidad', 'La interfaz debe ser intuitiva, sin necesidad de manual de usuario.'),
    ('RNF-03', 'Fiabilidad', 'El sistema debe manejar errores (JSON inválido, mensajes vacíos) sin caer.'),
    ('RNF-04', 'Mantenibilidad', 'El código debe seguir una arquitectura en capas para facilitar cambios.'),
    ('RNF-05', 'Portabilidad', 'El sistema debe ejecutarse en cualquier SO con Python 3.10+.'),
    ('RNF-06', 'Compatibilidad', 'La interfaz debe funcionar en navegadores modernos (Chrome, Firefox, Edge).'),
    ('RNF-07', 'Estética', 'El diseño visual debe seguir una temática coherente (filosofía china).'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = 'Característica'
hdr[2].text = 'Descripción'
for rid, rname, rdesc in req_nofunc:
    row = table.add_row().cells
    row[0].text = rid
    row[1].text = rname
    row[2].text = rdesc

doc.add_page_break()

# ============================================================
# 4. ARQUITECTURA DEL SISTEMA
# ============================================================
add_heading_apa('4. Arquitectura del Sistema', level=1)

add_heading_apa('4.1. Arquitectura en Capas', level=2)
add_normal_paragraph(
    'El sistema sigue una arquitectura en capas (layered architecture) que separa las '
    'responsabilidades en cinco niveles: Presentación, Aplicación, Lógica de Negocio, '
    'Inteligencia y Datos. Esta separación facilita el mantenimiento, las pruebas y '
    'la evolución del sistema.'
)

capas = [
    ('Capa de Presentación', 'index.html, style.css, chat.js', 'Interfaz de usuario con diseño chino, partículas, animaciones'),
    ('Capa de Aplicación', 'views.py, urls.py, models.py', 'Gestión HTTP, rutas y persistencia en Django'),
    ('Capa de Lógica', 'response_layer.py, calculator.py', 'Orquestación del flujo, detección de operaciones matemáticas'),
    ('Capa de Inteligencia', 'processing_layer.py, intelligence_layer.py', 'Normalización NLP, vectorización TF-IDF, red neuronal MLP'),
    ('Capa de Datos', 'data_layer.py, db.sqlite3', 'Base de conocimiento con 16 intenciones y almacenamiento'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Capa'
hdr[1].text = 'Componentes'
hdr[2].text = 'Responsabilidad'
for capa, comp, resp in capas:
    row = table.add_row().cells
    row[0].text = capa
    row[1].text = comp
    row[2].text = resp

add_heading_apa('4.2. Red Neuronal (MLPClassifier)', level=2)
add_normal_paragraph(
    'El modelo de inteligencia artificial utilizado es un Perceptrón Multicapa (MLP) '
    'implementado con scikit-learn. La red neuronal se entrena automáticamente al '
    'iniciar el servidor Django, utilizando los patrones definidos en la capa de datos.'
)

add_paragraph_bold('Arquitectura de la Red:')
nn_params = [
    ('Capa de Entrada', 'Vector TF-IDF', '500 características (max_features)'),
    ('Capa Oculta 1', '128 neuronas', 'Activación ReLU'),
    ('Capa Oculta 2', '64 neuronas', 'Activación ReLU'),
    ('Capa de Salida', '16 neuronas', 'Softmax (una por intención)'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Capa'
hdr[1].text = 'Configuración'
hdr[2].text = 'Detalle'
for capa, config, detalle in nn_params:
    row = table.add_row().cells
    row[0].text = capa
    row[1].text = config
    row[2].text = detalle

add_paragraph_bold('Hiperparámetros:')
hiper = [
    ('Optimizador', 'Adam'),
    ('Tasa de Aprendizaje', '0.01 (adaptativa)'),
    ('Regularización L2', '0.0001'),
    ('Máx. Iteraciones', '1000'),
    ('Early Stopping', 'Desactivado (dataset pequeño)'),
    ('Data Augmentation', '5 prefijos por patrón (x5 datos)'),
    ('Datos de Entrenamiento', '149 patrones base → 745 con augmentation'),
    ('Pérdida Final', '0.0106'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Parámetro'
hdr[1].text = 'Valor'
for param, val in hiper:
    row = table.add_row().cells
    row[0].text = param
    row[1].text = val

add_heading_apa('4.3. Módulo Calculadora', level=2)
add_normal_paragraph(
    'La calculadora es un módulo basado en expresiones regulares (regex) que intercepta '
    'mensajes antes de enviarlos a la red neuronal. Detecta cinco patrones de operaciones '
    'matemáticas en lenguaje natural y simbólico.'
)

patrones_calc = [
    ('Símbolos', '5 + 3, 10 - 4, 7 * 8, 100 / 5', r'\d+\s*[+\-*/]\s*\d+'),
    ('Verbos', 'suma 10 y 20, resta 50 y 15', r'(suma|resta|...) \d+ y \d+'),
    ('Palabras', '5 mas 3, 10 menos 2, 8 por 6', r'\d+ (mas|menos|por|entre) \d+'),
    ('Pregunta', 'cuanto es 9 + 11', r'cuanto es \d+ [op] \d+'),
    ('Sumale/Restale', 'sumale 5 a 10, restale 2 a 8', r'(sumale|restale) \d+ a \d+'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Patrón'
hdr[1].text = 'Ejemplo'
hdr[2].text = 'Regex'
for pat, ej, reg in patrones_calc:
    row = table.add_row().cells
    row[0].text = pat
    row[1].text = ej
    row[2].text = reg

doc.add_page_break()

# ============================================================
# 5. MODELO DE DATOS
# ============================================================
add_heading_apa('5. Modelo de Datos', level=1)
add_normal_paragraph(
    'El sistema utiliza SQLite como base de datos, con un único modelo Django para '
    'almacenar el historial de conversaciones.'
)

campos = [
    ('id', 'AutoField', 'PK', 'Identificador único auto-incremental'),
    ('user_message', 'TextField', 'NOT NULL', 'Mensaje escrito por el usuario'),
    ('bot_response', 'TextField', 'NOT NULL', 'Respuesta generada por el bot'),
    ('intent', 'CharField(100)', 'NOT NULL', 'Intención detectada por la RN'),
    ('confidence', 'FloatField', 'DEFAULT 0.0', 'Nivel de confianza (0.0 a 1.0)'),
    ('created_at', 'DateTimeField', 'AUTO', 'Timestamp de creación automático'),
]
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Campo'
hdr[1].text = 'Tipo Django'
hdr[2].text = 'Restricción'
hdr[3].text = 'Descripción'
for campo, tipo, rest, desc in campos:
    row = table.add_row().cells
    row[0].text = campo
    row[1].text = tipo
    row[2].text = rest
    row[3].text = desc

doc.add_page_break()

# ============================================================
# 6. INTERFAZ DE USUARIO
# ============================================================
add_heading_apa('6. Interfaz de Usuario', level=1)
add_normal_paragraph(
    'La interfaz del chatbot está inspirada en la filosofía y estética china, '
    'utilizando una paleta de colores elegante y elementos decorativos temáticos.'
)

add_paragraph_bold('Paleta de Colores:')
colores = [
    ('Negro Tinta', '#0a0a0f', 'Fondo principal'),
    ('Jade', '#2d6a4f', 'Elementos del bot, acentos'),
    ('Rojo Imperial', '#8B0000', 'Detalles y alertas'),
    ('Dorado', '#d4a574', 'Títulos y resaltados'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Nombre'
hdr[1].text = 'Código HEX'
hdr[2].text = 'Uso'
for nombre, hex_col, uso in colores:
    row = table.add_row().cells
    row[0].text = nombre
    row[1].text = hex_col
    row[2].text = uso

add_paragraph_bold('Elementos de la Interfaz:')
elementos = [
    'Header con avatar de dragón (龍), título LóngBot y estado en línea',
    'Barra de visualización de la arquitectura de la red neuronal',
    'Área de chat con burbujas de usuario y bot con timestamps',
    'Barra de intención con indicador de confianza visual',
    'Input con ornamento de bambú (竹) y botón de envío',
    'Footer con citas filosóficas rotativas',
    'Partículas de fondo animadas (luciérnagas)',
    'Botones de acciones rápidas en pantalla de bienvenida',
]
for elem in elementos:
    doc.add_paragraph(elem, style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. PLAN DE PRUEBAS
# ============================================================
add_heading_apa('7. Plan de Pruebas', level=1)
add_normal_paragraph(
    'Las pruebas del sistema verifican tanto la funcionalidad del motor NLP como '
    'la calculadora integrada y los endpoints de la API.'
)

pruebas = [
    ('TP-01', 'Enviar "hola"', 'Detectar intención "saludo" con confianza >= 95%', 'Aprobado'),
    ('TP-02', 'Enviar "que hora es"', 'Detectar intención "hora" y mostrar hora actual', 'Aprobado'),
    ('TP-03', 'Enviar "5 + 3"', 'Retornar resultado 8 con intención "calculadora_suma"', 'Aprobado'),
    ('TP-04', 'Enviar "multiplica 4 por 7"', 'Retornar resultado 28', 'Aprobado'),
    ('TP-05', 'Enviar "divide 100 entre 5"', 'Retornar resultado 20', 'Aprobado'),
    ('TP-06', 'Enviar "0 / 0"', 'Retornar error de división entre cero', 'Aprobado'),
    ('TP-07', 'Enviar mensaje vacío', 'Retornar error 400', 'Aprobado'),
    ('TP-08', 'Enviar "texto aleatorio xyz"', 'Retornar respuesta por defecto', 'Aprobado'),
    ('TP-09', 'Enviar JSON inválido', 'Retornar error 400 con mensaje descriptivo', 'Aprobado'),
    ('TP-10', 'Verificar persistencia', 'Comprobar que la conversación se guardó en BD', 'Aprobado'),
]
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = 'Caso de Prueba'
hdr[2].text = 'Resultado Esperado'
hdr[3].text = 'Estado'
for tid, caso, esperado, estado in pruebas:
    row = table.add_row().cells
    row[0].text = tid
    row[1].text = caso
    row[2].text = esperado
    row[3].text = estado

doc.add_page_break()

# ============================================================
# 8. UBICACIÓN DE FUNCIONES PRINCIPALES
# ============================================================
add_heading_apa('8. Ubicación de Funciones Principales de la Red Neuronal', level=1)
add_normal_paragraph(
    'A continuación se detallan las ubicaciones exactas de las funciones principales '
    'que conforman las dos redes/componentes de inteligencia del sistema: la Red Neuronal MLP '
    'y el Procesador de Texto (vectorizador TF-IDF).'
)

add_heading_apa('8.1. Red Neuronal MLP (intelligence_layer.py)', level=2)
add_normal_paragraph('Archivo: chatbot/engine/intelligence_layer.py')
funcs_rn = [
    ('RedNeuronal.__init__', 'Inicializa el MLPClassifier con capas (128, 64), ReLU, Adam, lr=0.01'),
    ('RedNeuronal.entrenar(X, etiquetas)', 'Entrena la red con vectores TF-IDF y etiquetas codificadas'),
    ('RedNeuronal.predecir(X)', 'Predice la intención y confianza usando predict_proba'),
    ('RedNeuronal.obtener_info_arquitectura()', 'Retorna info de la red: capas, intenciones, pérdida'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Función'
hdr[1].text = 'Descripción'
for func, desc in funcs_rn:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

add_heading_apa('8.2. Procesador de Texto TF-IDF (processing_layer.py)', level=2)
add_normal_paragraph('Archivo: chatbot/engine/processing_layer.py')
funcs_pt = [
    ('ProcesadorTexto.normalizar(texto)', 'Normaliza: minúsculas, quita acentos (unidecode), elimina puntuación'),
    ('ProcesadorTexto.entrenar_vectorizador(textos)', 'Entrena TfidfVectorizer con unigramas+bigramas, max 500 features'),
    ('ProcesadorTexto.vectorizar(textos)', 'Transforma lista de textos a matriz TF-IDF dispersa'),
    ('ProcesadorTexto.vectorizar_uno(texto)', 'Normaliza y vectoriza un solo texto para predicción'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Función'
hdr[1].text = 'Descripción'
for func, desc in funcs_pt:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

add_heading_apa('8.3. Motor Principal (response_layer.py)', level=2)
add_normal_paragraph('Archivo: chatbot/engine/response_layer.py')
funcs_motor = [
    ('MotorChatbot.entrenar()', 'Orquesta: obtener datos → augmentation → entrenar TF-IDF → entrenar MLP'),
    ('MotorChatbot.procesar_mensaje(mensaje)', 'Flujo: calculadora → vectorización → predicción → respuesta'),
    ('MotorChatbot.obtener_info()', 'Retorna información de la arquitectura para el frontend'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Función'
hdr[1].text = 'Descripción'
for func, desc in funcs_motor:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

add_heading_apa('8.4. Calculadora (calculator.py)', level=2)
add_normal_paragraph('Archivo: chatbot/engine/calculator.py')
funcs_calc = [
    ('detectar_operacion(texto)', 'Detecta y resuelve operaciones con 5 patrones regex'),
    ('es_operacion_matematica(texto)', 'Verificación rápida de si el texto contiene una operación'),
    ('_ejecutar_operacion(num1, op, num2)', 'Ejecuta +, -, *, / y maneja división entre cero'),
    ('_formatear_respuesta()', 'Genera respuesta elegante con el resultado y cita china'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Función'
hdr[1].text = 'Descripción'
for func, desc in funcs_calc:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

doc.add_page_break()

# ============================================================
# 9. REFERENCIAS BIBLIOGRÁFICAS (APA)
# ============================================================
add_heading_apa('9. Referencias Bibliográficas', level=1)
add_normal_paragraph(
    'Django Software Foundation. (2024). Django documentation (Version 5.2). '
    'https://docs.djangoproject.com/en/5.2/'
)
add_normal_paragraph(
    'IEEE Computer Society. (1998). IEEE Std 830-1998: IEEE recommended practice for '
    'software requirements specifications. IEEE.'
)
add_normal_paragraph(
    'ISO/IEC. (2011). ISO/IEC 25010:2011 — Systems and software engineering — Systems '
    'and software Quality Requirements and Evaluation (SQuaRE) — System and software '
    'quality models. International Organization for Standardization.'
)
add_normal_paragraph(
    'Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., '
    'Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., '
    'Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). Scikit-learn: '
    'Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.'
)
add_normal_paragraph(
    'Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text '
    'retrieval. Information Processing & Management, 24(5), 513–523. '
    'https://doi.org/10.1016/0306-4573(88)90021-0'
)

# ============================================================
# GUARDAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doc', 'Documentacion_LongBot_IEEE_APA.docx')
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
